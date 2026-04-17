# ============================================================
# DLSU CpE AI Academic Advising System
# Document Chunking Pipeline — LangChain + ChromaDB
# ============================================================
# Collections:
#   - checklist  : one chunk per course row (CPE_ID_118_CHECKLIST.pdf)
#   - policies   : section-aware chunks for all policy text, PDFs, and docx
#   - faqs       : one chunk per Q&A pair
#
# Ingestion functions:
#   chunk_and_ingest_checklist()      : parses checklist PDF → ChromaDB
#   chunk_and_ingest_policy()         : parses policy PDF or DOCX → ChromaDB
#   chunk_and_ingest_policy_text()    : ingests pre-extracted policy text → ChromaDB
#   chunk_and_ingest_faqs()           : ingests FAQ list → ChromaDB
# ============================================================

import uuid
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

import sys
sys.stdout.reconfigure(encoding='utf-8')

import pdfplumber                          # pip install pdfplumber
from docx import Document as DocxDocument  # pip install python-docx

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings


# ============================================================
# EMBEDDING MODEL
# intfloat/e5-small-v2 selected via embedding comparison experiment
# Requires "query: " prefix for queries, "passage: " for documents
# ============================================================

EMBEDDING_MODEL_NAME = "intfloat/e5-small-v2"
E5_DOC_PREFIX        = "passage: "   # prepended to all corpus documents
E5_QUERY_PREFIX      = "query: "     # prepended to all search queries

embedding_model = HuggingFaceEmbeddings(
    model_name=EMBEDDING_MODEL_NAME,
    model_kwargs={"device": "cpu"},       # switch to "cuda" if GPU available
    encode_kwargs={"normalize_embeddings": True},
)


# ============================================================
# CHROMADB COLLECTIONS
# ============================================================

CHROMA_BASE_DIR = "./chroma_store"

checklist_vectorstore = Chroma(
    collection_name="checklist",
    embedding_function=embedding_model,
    persist_directory=f"{CHROMA_BASE_DIR}/checklist",
)

policies_vectorstore = Chroma(
    collection_name="policies",
    embedding_function=embedding_model,
    persist_directory=f"{CHROMA_BASE_DIR}/policies",
)

faq_vectorstore = Chroma(
    collection_name="faqs",
    embedding_function=embedding_model,
    persist_directory=f"{CHROMA_BASE_DIR}/faqs",
)


# ============================================================
# SHARED POLICY SPLITTER
# Used by all three policy ingestion paths:
#   - chunk_and_ingest_policy()       (from file)
#   - chunk_and_ingest_policy_text()  (from pre-extracted text)
# ============================================================

POLICY_CHUNK_SIZE    = 256   # reduced from 512 to improve policy retrieval coverage
POLICY_CHUNK_OVERLAP = 50    # reduced proportionally

policy_splitter = RecursiveCharacterTextSplitter(
    chunk_size=POLICY_CHUNK_SIZE,
    chunk_overlap=POLICY_CHUNK_OVERLAP,
    separators=[
        "\n\n\n",   # major section breaks
        "\n\n",     # paragraph breaks
        "\n",       # line breaks
        ". ",       # sentence breaks
        " ",        # word breaks (last resort)
    ],
    length_function=len,
)

# Maps keyword tuples to doc_type — word-level matching handles
# spaces vs underscores in filenames (e.g. "GCOE UG OJT Policy (1).pdf")
POLICY_DOC_TYPE_KEYWORDS = [
    (["OJT", "Practicum", "Undergraduate"],          "ojt_policy"),
    (["Academic", "Advising", "Best", "Practices"],  "advising_best_practices"),
    (["Advising", "Guidelines"],                      "advising_guidelines"),
    (["Thesis", "Policies", "Guidelines"],            "thesis_policies"),
]


def resolve_doc_type(file_path: str) -> str:
    """
    Infer doc_type from filename using keyword matching.
    Handles spaces vs underscores in filenames.
    """
    words = set(re.sub(r"[_\-\.()]", " ", Path(file_path).name).upper().split())
    for keywords, doc_type in POLICY_DOC_TYPE_KEYWORDS:
        if all(k.upper() in words for k in keywords):
            return doc_type
    for keywords, doc_type in POLICY_DOC_TYPE_KEYWORDS:
        if any(k.upper() in words for k in keywords[:2]):
            return doc_type
    return "policy"


# ============================================================
# PARSER: CHECKLIST PDF
# File: CPE_ID_118_CHECKLIST.pdf
#
# Structure:
#   - Two-column table layout, 12 terms across 2 pages
#   - Columns: COURSE CODE | COURSE TITLE | UNITS | PREREQ TYPE | PREREQ CODES
#   - Prereq types: H (hard), S (soft), C (co-requisite)
#   - Multiple prereqs separated by ";" or "/"
#   - Term headers: "FIRST TERM" through "TWELFTH TERM"
#   - Year level inferred from term number (terms 1-3 = Y1, etc.)
# ============================================================

TERM_MAP = {
    "FIRST TERM":    (1, 1),  "SECOND TERM":  (1, 2),  "THIRD TERM":   (1, 3),
    "FOURTH TERM":   (2, 1),  "FIFTH TERM":   (2, 2),  "SIXTH TERM":   (2, 3),
    "SEVENTH TERM":  (3, 1),  "EIGHTH TERM":  (3, 2),  "NINTH TERM":   (3, 3),
    "TENTH TERM":    (4, 1),  "ELEVENTH TERM":(4, 2),  "TWELFTH TERM": (4, 3),
}

SKIP_PATTERNS = re.compile(
    r"^(COURSE|COURSE TITLE|UNITS|PREREQUISITES|TOTAL|LEGEND|"
    r"Please|This checklist|Chair|GCOE|Prepared|Approved|Noted).*",
    re.IGNORECASE
)


def parse_prerequisite_string(raw: str) -> List[Dict[str, str]]:
    """
    Parse raw prerequisite string into structured list.

    Examples:
        "H CALENG1"            -> [{"type": "H", "course_code": "CALENG1"}]
        "S/H CALENG1/BASPHYS"  -> [{"type": "S", "course_code": "CALENG1"},
                                    {"type": "H", "course_code": "BASPHYS"}]
        "H FDCNSYS; MICPROS"   -> [{"type": "H", "course_code": "FDCNSYS"},
                                    {"type": "H", "course_code": "MICPROS"}]
        "3rd Yr Standing"      -> [{"type": "standing", "course_code": "3rd Yr Standing"}]
        ""                     -> []
    """
    if not raw or raw.strip() == "":
        return []

    raw = raw.strip()

    if "standing" in raw.lower():
        return [{"type": "standing", "course_code": raw}]

    prerequisites = []
    groups = [g.strip() for g in raw.split(";")]

    for group in groups:
        type_match = re.match(r"^([A-Z](?:/[A-Z])*)\s+(.+)$", group)
        if type_match:
            types = type_match.group(1).split("/")
            codes = [c.strip() for c in type_match.group(2).split("/")]
            for i, code in enumerate(codes):
                req_type = types[i] if i < len(types) else types[-1]
                if code:
                    prerequisites.append({"type": req_type, "course_code": code})
        else:
            for code in group.split("/"):
                code = code.strip()
                if code:
                    prerequisites.append({"type": "unknown", "course_code": code})

    return prerequisites


def _extract_course_from_cells(
    cells: List[str],
    start: int,
    term_name: Optional[str],
    year_level: Optional[int],
    term_number: Optional[int],
    curriculum_year: str,
) -> Optional[Dict[str, Any]]:
    """Extract one course entry from a slice of table cells."""
    try:
        code      = cells[start].strip()     if start     < len(cells) else ""
        title     = cells[start+1].strip()   if start+1   < len(cells) else ""
        units_raw = cells[start+2].strip()   if start+2   < len(cells) else ""
        p_type    = cells[start+3].strip()   if start+3   < len(cells) else ""
        p_codes   = cells[start+4].strip()   if start+4   < len(cells) else ""
    except IndexError:
        return None

    if not re.match(r"^[A-Z]{3,8}\d*[A-Z]?$", code):
        return None

    units = 0
    m = re.search(r"\d+", units_raw)
    if m:
        units = int(m.group())

    prerequisites = parse_prerequisite_string(f"{p_type} {p_codes}".strip())

    return {
        "course_code":     code,
        "title":           title,
        "units":           units,
        "year_level":      year_level,
        "term_number":     term_number,
        "term_name":       term_name,
        "prerequisites":   prerequisites,
        "curriculum_year": curriculum_year,
    }


def parse_checklist_pdf(file_path: str, curriculum_year: str = "2018") -> List[Dict[str, Any]]:
    """
    Parse the DLSU CpE course checklist PDF into a list of course dicts.

    Args:
        file_path:       Path to CPE_ID_118_CHECKLIST.pdf.
        curriculum_year: AY the checklist was introduced.

    Returns:
        List of course dicts with keys:
        course_code, title, units, year_level, term_number,
        term_name, prerequisites, curriculum_year.
    """
    courses = []
    current_term = current_year_level = current_term_number = None

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            for table in (page.extract_tables() or []):
                for row in table:
                    cleaned = [str(c).strip() if c else "" for c in row]
                    if all(c == "" for c in cleaned):
                        continue

                    full_row = " ".join(c for c in cleaned if c).upper()
                    for term_name, (yr, tm) in TERM_MAP.items():
                        if term_name in full_row:
                            current_term = term_name
                            current_year_level = yr
                            current_term_number = tm
                            break

                    if SKIP_PATTERNS.match(cleaned[0] if cleaned else ""):
                        continue

                    left = _extract_course_from_cells(
                        cleaned, 0,
                        current_term, current_year_level,
                        current_term_number, curriculum_year,
                    )
                    if left:
                        courses.append(left)

                    if len(cleaned) > 4:
                        right = _extract_course_from_cells(
                            cleaned, 4,
                            current_term, current_year_level,
                            current_term_number, curriculum_year,
                        )
                        if right:
                            courses.append(right)

    print(f"[Checklist] Parsed {len(courses)} courses from '{file_path}'.")
    return courses


# ============================================================
# PARSER: POLICY PDF
# ============================================================

def parse_policy_pdf(file_path: str) -> str:
    """Extract all text from a policy PDF file."""
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
    full_text = "\n\n".join(parts)
    print(f"[Policy PDF] Extracted {len(full_text)} chars from '{file_path}'.")
    return full_text


# ============================================================
# PARSER: POLICY DOCX
# ============================================================

def parse_policy_docx(file_path: str) -> str:
    """
    Extract all text from a policy .docx file.
    Preserves heading structure and extracts table cell text.
    """
    doc = DocxDocument(file_path)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            for p in doc.paragraphs:
                if p._element is element:
                    text = p.text.strip()
                    if text:
                        parts.append(f"\n\n{text}" if p.style.name.startswith("Heading") else text)
                    break

        elif tag == "tbl":
            NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for row in element.findall(f".//{{{NS}}}tr"):
                row_texts = []
                for cell in row.findall(f"{{{NS}}}tc"):
                    cell_text = " ".join(
                        n.text for n in cell.iter() if n.text and n.text.strip()
                    ).strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

    full_text = "\n".join(parts)
    print(f"[Policy DOCX] Extracted {len(full_text)} chars from '{file_path}'.")
    return full_text


# ============================================================
# INGESTION: CHECKLIST
# ============================================================

def format_course_chunk(row: Dict[str, Any]) -> str:
    """Convert a course row dict into a readable embeddable text chunk."""
    prereqs = row.get("prerequisites", [])
    prereq_str = (
        "\n".join(f"  - {p['course_code']} ({p['type']})" for p in prereqs)
        if prereqs else "  - None"
    )
    return (
        f"Course Code: {row['course_code']}\n"
        f"Title: {row['title']}\n"
        f"Units: {row['units']}\n"
        f"Year Level: {row.get('year_level', 'N/A')} | "
        f"Term: {row.get('term_number', 'N/A')} ({row.get('term_name', 'N/A')})\n"
        f"Curriculum Year: {row.get('curriculum_year', 'N/A')}\n"
        f"Prerequisites:\n{prereq_str}"
    )


def chunk_and_ingest_checklist(file_path: str, curriculum_year: str = "2018"):
    """Parse the checklist PDF and store one chunk per course into ChromaDB."""
    rows = parse_checklist_pdf(file_path, curriculum_year)
    documents = []

    for i, row in enumerate(rows):
        prereq_codes = [p["course_code"] for p in row.get("prerequisites", [])]
        documents.append(Document(
            page_content=format_course_chunk(row),
            metadata={
                "doc_id":          str(uuid.uuid4()),
                "source":          Path(file_path).name,
                "doc_type":        "checklist",
                "curriculum_year": curriculum_year,
                "course_code":     row["course_code"],
                "year_level":      row.get("year_level"),
                "term_number":     row.get("term_number"),
                "term_name":       row.get("term_name", ""),
                "units":           row.get("units", 0),
                "prereq_codes":    ",".join(prereq_codes),
                "model_name":      EMBEDDING_MODEL_NAME,
                "chunk_index":     i,
                "chunk_total":     len(rows),
            }
        ))

    checklist_vectorstore.add_documents(documents)
    
    print(f"[Checklist] {len(documents)} chunks stored in ChromaDB.")


# ============================================================
# INGESTION: POLICY FROM FILE (PDF or DOCX)
# ============================================================

def chunk_and_ingest_policy(
    file_path: str,
    title: Optional[str] = None,
    version: str = "2024",
):
    """
    Parse a policy PDF or DOCX file and store chunks into ChromaDB.
    File type is detected automatically from the file extension.
    """
    path = Path(file_path)
    title = title or path.stem.replace("_", " ").replace("-", " ")
    doc_type = resolve_doc_type(file_path)

    if path.suffix.lower() == ".pdf":
        raw_text = parse_policy_pdf(file_path)
    elif path.suffix.lower() in (".docx", ".doc"):
        raw_text = parse_policy_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    _ingest_policy_text_internal(raw_text, title, doc_type, path.name, version)


# ============================================================
# INGESTION: POLICY FROM PRE-EXTRACTED TEXT
# Used for handbook sections already stored in source_data.py
# ============================================================

def chunk_and_ingest_policy_text(
    text: str,
    title: str,
    doc_type: str,
    version: str = "Current",
    source: str = "DLSU Student Handbook",
):
    """
    Chunk and store pre-extracted policy text into the policies ChromaDB collection.

    Use this for handbook sections that are already available as plain text
    in source_data.py (e.g. retention policy, load policy, lab/lecture policy).

    Args:
        text:     The raw policy text to chunk and embed.
        title:    Human-readable title of the policy section.
        doc_type: Metadata tag for filtered retrieval
                  (e.g. 'retention_policy', 'load_policy', 'lab_lecture_policy').
        version:  Document version or effectivity date.
        source:   Origin of the document (e.g. 'DLSU Student Handbook').
    """
    _ingest_policy_text_internal(text, title, doc_type, source, version)


def _ingest_policy_text_internal(
    raw_text: str,
    title: str,
    doc_type: str,
    source: str,
    version: str,
):
    """Shared internal function: split text and store chunks in ChromaDB."""
    raw_chunks = policy_splitter.split_text(raw_text)
    documents = []

    for i, chunk in enumerate(raw_chunks):
        documents.append(Document(
            page_content=chunk,
            metadata={
                "doc_id":      str(uuid.uuid4()),
                "source":      source,
                "title":       title,
                "doc_type":    doc_type,
                "version":     version,
                "chunk_index": i,
                "chunk_total": len(raw_chunks),
                "model_name":  EMBEDDING_MODEL_NAME,
            }
        ))

    policies_vectorstore.add_documents(documents)
    
    print(
        f"[Policy] '{title}' → {len(raw_chunks)} chunks "
        f"(doc_type='{doc_type}') stored in ChromaDB."
    )


# ============================================================
# INGESTION: FAQs
# ============================================================

def chunk_and_ingest_faqs(faq_list: List[Dict[str, Any]]):
    """
    Ingest a list of FAQ dicts into ChromaDB.

    Each dict must have: question, answer, category, program,
    difficulty, verified. Unverified entries are skipped.
    """
    documents = []
    skipped = 0

    for i, faq in enumerate(faq_list):
        if not faq.get("verified", True):
            skipped += 1
            continue

        content = (
            f"Question: {faq.get('question', '').strip()}\n"
            f"Answer: {faq.get('answer', '').strip()}"
        )
        documents.append(Document(
            page_content=content,
            metadata={
                "doc_id":      str(uuid.uuid4()),
                "source":      "DLSU CpE Curated FAQs",
                "doc_type":    "faq",
                "category":    faq.get("category", "General"),
                "program":     faq.get("program", "BS Computer Engineering"),
                "difficulty":  faq.get("difficulty", "basic"),
                "verified":    True,
                "chunk_index": i,
                "chunk_total": len(faq_list),
                "model_name":  EMBEDDING_MODEL_NAME,
            }
        ))

    faq_vectorstore.add_documents(documents)
    
    print(f"[FAQs] {len(documents)} chunks stored. {skipped} unverified skipped.")


# ============================================================
# RETRIEVAL HELPERS
# ============================================================

def retrieve_checklist(query: str, k: int = 5) -> List[Document]:
    """Semantic search over the checklist collection."""
    return checklist_vectorstore.similarity_search(query, k=k)


def retrieve_policies(
    query: str,
    k: int = 5,
    doc_type: Optional[str] = None,
) -> List[Document]:
    """
    Semantic search over the policies collection.
    Optionally filter by doc_type, e.g.:
        'retention_policy', 'load_policy', 'lab_lecture_policy',
        'ojt_policy', 'advising_guidelines', 'thesis_policies',
        'crediting_process', 'academic_resources'
    """
    filter_dict = {"doc_type": doc_type} if doc_type else None
    return policies_vectorstore.similarity_search(query, k=k, filter=filter_dict)


def retrieve_faqs(
    query: str,
    k: int = 5,
    category: Optional[str] = None,
) -> List[Document]:
    """
    Semantic search over the FAQ collection.
    Optionally filter by category, e.g.:
        'Retention', 'Enrollment', 'Prerequisites',
        'Crediting', 'Resources', 'OJT', 'Advising'
    """
    filter_dict = {"category": category} if category else None
    return faq_vectorstore.similarity_search(query, k=k, filter=filter_dict)


def retrieve_all(query: str, k: int = 3) -> List[Document]:
    """Search all three collections and merge results (up to 3k total)."""
    return (
        retrieve_checklist(query, k=k)
        + retrieve_policies(query, k=k)
        + retrieve_faqs(query, k=k)
    )


# ============================================================
# MAIN — batch ingestion from parsed_data/ JSON files
# ============================================================
# Does NOT re-parse PDFs. Reads from:
#   parsed_data/checklist_rows.json   (from batch_parser.py)
#   parsed_data/policy_sections.json  (from batch_parser.py)
#   advising_dataset.xlsx             (593 Q&A pairs)
#   source_data.FAQ_LIST              (13 hand-written FAQs)
#
# Run:
#   python chunking_pipeline.py
#   python chunking_pipeline.py --parsed-dir ./parsed_data
# ============================================================

if __name__ == "__main__":
    import json
    import argparse
    import pandas as pd

    parser = argparse.ArgumentParser(
        description="Ingest all parsed documents into ChromaDB."
    )
    parser.add_argument("--parsed-dir", type=Path, default=Path("./parsed_data"))
    parser.add_argument("--dataset",    type=Path, default=Path("./dataset_train.xlsx"))
    args = parser.parse_args()

    # ----------------------------------------------------------
    # 1. Ingest all checklist courses from parsed JSON
    # ----------------------------------------------------------
    checklist_path = args.parsed_dir / "checklist_rows.json"
    if checklist_path.exists():
        rows = json.load(open(checklist_path, encoding="utf-8"))

        # Deduplicate by course_code + program + academic_year
        seen     = set()
        unique   = []
        for row in rows:
            key = f"{row['program']}_{row['course_code']}_{row.get('academic_year','')}"
            if key not in seen:
                seen.add(key)
                unique.append(row)

        print(f"\n[Checklist] Ingesting {len(unique)} unique course entries...")
        documents = []
        for i, row in enumerate(unique):
            prereqs = row.get("prerequisites", [])
            prereq_str = (
                "\n".join(f"  - {p['course_code']} ({p['type']})" for p in prereqs)
                if prereqs else "  - None"
            )
            # e5 models require "passage: " prefix on corpus documents
            content = (
                f"{E5_DOC_PREFIX}"
                f"Course Code: {row['course_code']}\n"
                f"Title: {row.get('title', '')}\n"
                f"Units: {row.get('units', '')}\n"
                f"Year Level: {row.get('year_level', 'N/A')} | "
                f"Term: {row.get('term_number', 'N/A')} ({row.get('term_name', 'N/A')})\n"
                f"Program: {row.get('program', 'N/A')} | "
                f"Curriculum Year: {row.get('academic_year', 'N/A')}\n"
                f"Prerequisites:\n{prereq_str}"
            )
            prereq_codes = [p["course_code"] for p in prereqs]
            documents.append(Document(
                page_content=content,
                metadata={
                    "doc_id":          str(uuid.uuid4()),
                    "source":          f"{row.get('program','')} Checklist",
                    "doc_type":        "checklist",
                    "curriculum_year": row.get("academic_year", ""),
                    "course_code":     row["course_code"],
                    "year_level":      row.get("year_level"),
                    "term_number":     row.get("term_number"),
                    "term_name":       row.get("term_name", ""),
                    "units":           row.get("units", 0),
                    "program":         row.get("program", ""),
                    "prereq_codes":    ",".join(prereq_codes),
                    "model_name":      EMBEDDING_MODEL_NAME,
                    "chunk_index":     i,
                    "chunk_total":     len(unique),
                }
            ))
            # Batch ingest every 200 docs
            if len(documents) >= 200:
                checklist_vectorstore.add_documents(documents)
                documents = []

        if documents:
            checklist_vectorstore.add_documents(documents)
        
        print(f"[Checklist] {len(unique)} course chunks stored in ChromaDB.")
    else:
        print(f"[WARN] {checklist_path} not found — run batch_parser.py first.")

    # ----------------------------------------------------------
    # 2. Ingest all policy sections from parsed JSON
    # ----------------------------------------------------------
    policy_path = args.parsed_dir / "policy_sections.json"
    if policy_path.exists():
        sections = json.load(open(policy_path, encoding="utf-8"))
        print(f"\n[Policy] Ingesting {len(sections)} policy section(s)...")
        for section in sections:
            chunk_and_ingest_policy_text(
                text=f"{section.get('text', '')}",
                title=section.get("title", ""),
                doc_type=section.get("doc_type", "policy"),
                version=section.get("version", "2024"),
                source=section.get("source", "DLSU GCOE"),
            )
    else:
        print(f"[WARN] {policy_path} not found — run batch_parser.py first.")

    # ----------------------------------------------------------
    # 3. Ingest FAQs from dataset_train.xlsx (train split only)
    # ----------------------------------------------------------
    if args.dataset.exists():
        df = pd.read_excel(args.dataset)
        # Normalise column names — robust to old and new dataset layouts
        df.columns = [str(c).strip() for c in df.columns]
        cat_col  = "category"        if "category"        in df.columns else None
        prog_col = "program"         if "program"         in df.columns else None
        faq_list = []
        for _, row in df.iterrows():
            q = row.get("question", "")
            a = row.get("answer", "")
            if not q or not a or str(q).strip() == "nan" or str(a).strip() == "nan":
                continue
            faq_list.append({
                "question": str(q).strip(),
                "answer":   str(a).strip(),
                "category": str(row[cat_col])  if cat_col  else "General",
                "program":  str(row[prog_col]) if prog_col else "GENERAL",
                "difficulty": "basic",
                "verified":   True,
            })
        print(f"\n[FAQs] Ingesting {len(faq_list)} Q&A pairs from {args.dataset.name}...")
        chunk_and_ingest_faqs(faq_list)
    else:
        print(f"[WARN] {args.dataset} not found — skipping dataset FAQs.")

    # ----------------------------------------------------------
    # 4. Ingest hand-written FAQs from source_data.py
    # ----------------------------------------------------------
    try:
        from source_data import FAQ_LIST
        verified = [f for f in FAQ_LIST if f.get("verified", True)]
        print(f"\n[FAQs] Ingesting {len(verified)} hand-written FAQs from source_data...")
        chunk_and_ingest_faqs(verified)
    except ImportError:
        print("[WARN] source_data.py not found — skipping hand-written FAQs.")

    # ----------------------------------------------------------
    # 5. Test retrieval
    # ----------------------------------------------------------
    print("\n" + "="*55)
    print("TEST RETRIEVAL")
    print("="*55)
    test_queries = [
        "How many units of failure will make me ineligible?",
        "What is the maximum load per term?",
        "What are the prerequisites for CONETSC?",
        "How many hours do I need for OJT?",
        "Do I need to retake both lecture and lab if I fail one?",
    ]
    for q in test_queries:
        print(f"\nQuery: {q}")
        # e5 queries need the query prefix
        results = retrieve_all(f"{E5_QUERY_PREFIX}{q}", k=2)
        for r in results:
            print(f"  [{r.metadata.get('doc_type')}] "
                  f"{r.page_content[:100].strip().replace(chr(10), ' ')}...")
